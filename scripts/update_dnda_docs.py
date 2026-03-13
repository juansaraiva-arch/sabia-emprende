#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
update_dnda_docs.py
-------------------
Actualiza los documentos DNDA (Manual de Uso y Memoria Descriptiva)
con los nuevos modulos v1.1.0, y genera el HTML de Codigo Fuente.

Requisitos: python-docx >= 1.2.0
Uso: python scripts/update_dnda_docs.py
"""

import os
import sys
import html as html_module
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ───────────────── PATHS ─────────────────

REPO_ROOT = Path(r"C:\Users\FranciscoSaraiva\sabia-emprende")
DOWNLOADS = Path(r"C:\Users\FranciscoSaraiva\Downloads")

MANUAL_IN = DOWNLOADS / "DNDA_ManualDeUso_MiDirectorFinancieroPTY_v1.docx"
MANUAL_OUT = DOWNLOADS / "DNDA_ManualDeUso_MiDirectorFinancieroPTY_v2.docx"

MEMORIA_IN = DOWNLOADS / "DNDA_MemoriaDescriptiva_MiDirectorFinancieroPTY_v1.docx"
MEMORIA_OUT = DOWNLOADS / "DNDA_MemoriaDescriptiva_MiDirectorFinancieroPTY_v2.docx"

HTML_OUT = REPO_ROOT / "docs" / "dnda" / "DNDA_CodigoFuente_PrimerasUltimas.html"

# ───────────────── HELPERS ─────────────────

def find_paragraph(doc, search_text, start_from=0):
    """Encuentra el primer parrafo que contenga search_text (desde start_from)."""
    for i, p in enumerate(doc.paragraphs):
        if i < start_from:
            continue
        if search_text in p.text:
            return i, p
    return None, None


def find_all_paragraphs(doc, search_text):
    """Encuentra todos los parrafos que contengan search_text."""
    results = []
    for i, p in enumerate(doc.paragraphs):
        if search_text in p.text:
            results.append((i, p))
    return results


def insert_paragraph_after(paragraph, text, style_name=None):
    """Inserta un nuevo parrafo DESPUES del parrafo dado usando manipulacion XML."""
    new_p_element = OxmlElement('w:p')
    paragraph._p.addnext(new_p_element)
    new_para = Paragraph(new_p_element, paragraph._parent)
    if style_name:
        try:
            new_para.style = style_name
        except Exception:
            pass
    if text:
        new_para.text = text
    return new_para


def insert_multiple_after(paragraph, items):
    """
    Inserta multiples parrafos despues de paragraph.
    items: lista de (text, style_name) tuples
    Retorna el ultimo parrafo insertado.
    """
    current = paragraph
    for text, style in items:
        current = insert_paragraph_after(current, text, style)
    return current


def replace_text_in_all(doc, old_text, new_text):
    """Reemplaza texto en todos los parrafos del documento."""
    count = 0
    for p in doc.paragraphs:
        if old_text in p.text:
            for run in p.runs:
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)
                    count += 1
    # Tambien en tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if old_text in p.text:
                        for run in p.runs:
                            if old_text in run.text:
                                run.text = run.text.replace(old_text, new_text)
                                count += 1
    return count


# ═══════════════════════════════════════════════════════════════
#  TASK 1: UPDATE MANUAL DE USO
# ═══════════════════════════════════════════════════════════════

def update_manual():
    print("=" * 60)
    print("TASK 1: Actualizando Manual de Uso...")
    print("=" * 60)

    doc = Document(str(MANUAL_IN))

    # 1) Version 1.0.0 -> 1.1.0
    n = replace_text_in_all(doc, "1.0.0", "1.1.0")
    print(f"  Version: reemplazados {n} ocurrencias de 1.0.0 -> 1.1.0")

    # 2) Update TOC: add new entries after existing ones
    # Add after "5.5  Mi Inventario"
    idx, p = find_paragraph(doc, "5.5  Mi Inventario")
    if p:
        toc_new_5 = [
            ("5.6  Libro de Ventas", "Normal"),
            ("5.7  Cierre Mensual y Reportes", "Normal"),
        ]
        last = insert_multiple_after(p, toc_new_5)
        print(f"  TOC: insertadas entradas 5.6, 5.7 despues de indice {idx}")
    else:
        print("  WARN: No se encontro '5.5  Mi Inventario' en TOC")

    # Add after "6.5  Mas Herramientas Financieras"
    idx, p = find_paragraph(doc, "6.5  M")
    if p:
        toc_new_6 = [
            ("6.6  Forecast P&L 12 Meses", "Normal"),
            ("6.7  Simulador de N\u00f3mina con Desglose", "Normal"),
            ("6.8  Diagn\u00f3stico CB Insights", "Normal"),
        ]
        last = insert_multiple_after(p, toc_new_6)
        print(f"  TOC: insertadas entradas 6.6, 6.7, 6.8 despues de indice {idx}")

    # Add after "7.4  MUPA"
    idx, p = find_paragraph(doc, "7.4  MUPA")
    if p:
        toc_new_7 = [
            ("7.5  Traductor de Jerga Legal", "Normal"),
        ]
        last = insert_multiple_after(p, toc_new_7)
        print(f"  TOC: insertada entrada 7.5 despues de indice {idx}")

    # Add after "9  Sistema de Alertas" in TOC
    idx, p = find_paragraph(doc, "9  Sistema de Alertas")
    if p:
        toc_new_9 = [
            ("9.1  Alertas de Facturaci\u00f3n", "Normal"),
        ]
        last = insert_multiple_after(p, toc_new_9)
        print(f"  TOC: insertada entrada 9.1 despues de indice {idx}")

    # Add "10A  Onboarding Guiado" before "10  Preguntas Frecuentes" in TOC
    idx, p = find_paragraph(doc, "10  Preguntas Frecuentes")
    if p:
        # Insert BEFORE p by inserting after the previous paragraph
        # We need the previous paragraph
        prev_p = doc.paragraphs[idx - 1] if idx > 0 else None
        if prev_p:
            insert_paragraph_after(prev_p, "10A  Onboarding Guiado", "Normal")
            print(f"  TOC: insertada entrada 10A antes de Preguntas Frecuentes")

    # 3) Update "Mandibulas" reference in propuesta de valor
    for p in doc.paragraphs:
        if "Mand\u00edbulas" in p.text:
            for run in p.runs:
                if "Mand\u00edbulas" in run.text:
                    run.text = run.text.replace("Mand\u00edbulas", "Brecha de Rentabilidad")
            print(f"  Renombrado: Mandibulas -> Brecha de Rentabilidad en propuesta de valor")

    # 4) Add new capability items in section 1.3 (propuesta de valor)
    # Find last list item in propuesta de valor section (the alertas one)
    idx, p = find_paragraph(doc, "Sistema de alertas estrat\u00e9gicas con 7 categor\u00edas")
    if p:
        new_capabilities = [
            ("Libro de Ventas con 3 segmentos: venta r\u00e1pida manual, importaci\u00f3n CSV desde SFEP DGI, e integraci\u00f3n PAC Gosocket para facturaci\u00f3n electr\u00f3nica", "List Paragraph"),
            ("Forecast P&L a 12 meses con regresi\u00f3n lineal, supuestos editables y gr\u00e1ficos de proyecci\u00f3n", "List Paragraph"),
            ("Simulador de N\u00f3mina con desglose completo de 14 l\u00edneas empleador y c\u00e1lculo ISR paso a paso", "List Paragraph"),
            ("Diagn\u00f3stico CB Insights: detector de las 6 causas principales de fracaso empresarial con score 0-100", "List Paragraph"),
            ("Traductor de Jerga Legal: traducci\u00f3n de documentos legales a lenguaje sencillo usando GPT-4o", "List Paragraph"),
            ("Cierre Mensual con generaci\u00f3n de reportes autom\u00e1ticos y snapshots de KPIs", "List Paragraph"),
            ("Onboarding Guiado con tour modal de 5 pasos y tooltips flotantes (Silver Economy friendly)", "List Paragraph"),
            ("Panel de m\u00e9tricas admin con MAU, churn rate, DAU/MAU y cohortes de retenci\u00f3n", "List Paragraph"),
        ]
        insert_multiple_after(p, new_capabilities)
        print(f"  Propuesta de valor: agregadas {len(new_capabilities)} nuevas capacidades")

    # 5) Insert new section 5.6 Libro de Ventas (before Glosario)
    # Find "5.5 Mi Inventario" heading section end — we'll insert after the last paragraph of 5.5
    # Better: find the heading for section 6 (Mis Finanzas) and insert before it
    # Actually, insert AFTER the 5.5 Mi Inventario section content, before 6. Modulo Mis Finanzas
    idx_6, p_6 = find_paragraph(doc, "6. M\u00f3dulo Mis Finanzas")
    if not p_6:
        idx_6, p_6 = find_paragraph(doc, "M\u00f3dulo Mis Finanzas")

    if p_6:
        # Insert before section 6 by inserting after the previous paragraph
        prev_p = doc.paragraphs[idx_6 - 1]
        new_56_content = [
            ("5.6 Libro de Ventas", "Heading 2"),
            ("El Libro de Ventas permite registrar y gestionar todas las transacciones de venta del negocio. Opera en tres segmentos progresivos seg\u00fan el volumen del contribuyente, conforme a los umbrales establecidos por la DGI en la Resoluci\u00f3n 201-6299.", "Normal"),
            ("Segmento 1: Venta R\u00e1pida Manual", "Heading 3"),
            ("El primer segmento provee un bot\u00f3n de venta r\u00e1pida que permite registrar transacciones individuales con c\u00e1lculo autom\u00e1tico del ITBMS (7%). El usuario ingresa el monto, la descripci\u00f3n y la fecha, y el sistema genera autom\u00e1ticamente el desglose fiscal. Ideal para negocios con bajo volumen de facturaci\u00f3n.", "Normal"),
            ("Segmento 2: Importaci\u00f3n CSV desde SFEP DGI", "Heading 3"),
            ("El segundo segmento permite importar ventas masivas desde archivos CSV generados por el Sistema de Facturaci\u00f3n Electr\u00f3nica de Panam\u00e1 (SFEP) de la DGI. El parser reconoce autom\u00e1ticamente las columnas del formato oficial, validando RUC, montos e ITBMS. Se identifica con un badge azul en la tabla unificada.", "Normal"),
            ("Segmento 3: Integraci\u00f3n PAC Gosocket", "Heading 3"),
            ("El tercer segmento conecta con Gosocket, Proveedor Autorizado de Certificaci\u00f3n (PAC), para facturaci\u00f3n electr\u00f3nica obligatoria cuando el contribuyente supera los umbrales DGI: 100 facturas/mes o B/.36,000/a\u00f1o en ventas gravadas. Las facturas electr\u00f3nicas se identifican con badge verde en la tabla.", "Normal"),
            ("Tabla Unificada de Ventas", "Heading 3"),
            ("Todas las ventas, independientemente de su origen, se consolidan en una tabla unificada con badges visuales por origen: Manual (gris), DGI (azul) y PAC (verde). La tabla permite filtrado, ordenamiento y exportaci\u00f3n.", "Normal"),
            ("5.7 Cierre Mensual y Reportes", "Heading 2"),
            ("El m\u00f3dulo de Cierre Mensual genera reportes integrales del per\u00edodo que incluyen snapshots de KPIs financieros, cascada P&L del mes, resumen de ventas por segmento, estado de alertas, y comparativo con el mes anterior. Una vez cerrado un per\u00edodo, el sistema protege los registros contra re-apertura accidental, garantizando la integridad del historial contable.", "Normal"),
        ]
        insert_multiple_after(prev_p, new_56_content)
        print(f"  Secciones 5.6 y 5.7 insertadas antes de Modulo Mis Finanzas")

    # 6) Insert new sections 6.6, 6.7, 6.8 after section 6.5
    # Find "6.5.7 Contactos, Presupuestos y Tendencias"
    idx_657, p_657 = find_paragraph(doc, "6.5.7 Contactos, Presupuestos y Tendencias")
    if p_657:
        # Find end of 6.5.7 content (next paragraph after the description)
        # The description is the paragraph right after 6.5.7
        next_idx = idx_657 + 1
        while next_idx < len(doc.paragraphs):
            next_p = doc.paragraphs[next_idx]
            next_text = next_p.text.strip()
            if next_text and (next_p.style.name.startswith('Heading') or next_text.startswith('7.')):
                break
            next_idx += 1
        # Insert before the heading 7 section
        anchor = doc.paragraphs[next_idx - 1]
        new_66_content = [
            ("6.6 Forecast P&L 12 Meses", "Heading 2"),
            ("El Forecast P&L proyecta los resultados financieros del negocio a 12 meses utilizando regresi\u00f3n lineal client-side sobre los datos hist\u00f3ricos ingresados. El usuario puede ajustar cuatro supuestos editables: tasa de crecimiento de ingresos, incremento esperado de costos, inflaci\u00f3n salarial y variaci\u00f3n de gastos operativos.", "Normal"),
            ("La proyecci\u00f3n se presenta en una tabla P&L interactiva click-to-edit, donde cada celda puede modificarse manualmente para modelar escenarios espec\u00edficos. Un gr\u00e1fico de l\u00edneas (Recharts LineChart) visualiza la tendencia proyectada.", "Normal"),
            ("Los KPIs principales del forecast incluyen: revenue anual proyectado, EBITDA estimado, margen neto esperado y meses de runway con la tendencia actual.", "Normal"),
            ("6.7 Simulador de N\u00f3mina con Desglose Completo", "Heading 2"),
            ("El simulador de n\u00f3mina avanzado presenta el costo total de contratar empleados en Panam\u00e1 con un desglose exhaustivo de 14 l\u00edneas del lado empleador: CSS Patronal (12.25%), Seguro Educativo Patronal (1.50%), Riesgo Profesional (1.50%), D\u00e9cimo Tercer Mes (8.33%), Vacaciones (3.33%), Prima de Antig\u00fcedad (1.92%), subtotales parciales, y el Factor 1.36x multiplicador.", "Normal"),
            ("Del lado del empleado, el desglose incluye el c\u00e1lculo paso a paso del ISR: (1) salario bruto, (2) deducci\u00f3n de CSS obrera 9.75% + SE 1.25% = 11%, (3) base gravable mensual, (4) anualizaci\u00f3n \u00d712, (5) aplicaci\u00f3n de tabla progresiva DGI, (6) divisi\u00f3n /12 para obtener retenci\u00f3n mensual. Ejemplo verificado: B/.1,500 bruto \u2192 B/.62.75/mes de ISR.", "Normal"),
            ("La vista consolidada multi-empleado muestra el total de salida real de caja diferenciado de las provisiones mensuales (XIII Mes, vacaciones, antig\u00fcedad).", "Normal"),
            ("6.8 Diagn\u00f3stico CB Insights", "Heading 2"),
            ("El diagn\u00f3stico CB Insights implementa un detector client-side que eval\u00faa las 6 causas principales de fracaso empresarial seg\u00fan el estudio de CB Insights: (1) Sin mercado/demanda, (2) Sin caja/liquidez, (3) Equipo incorrecto, (4) Competencia superior, (5) Precios inadecuados, (6) Modelo no escalable.", "Normal"),
            ("Cada causa se eval\u00faa mediante algoritmos que analizan los datos financieros del usuario. El resultado se presenta como un widget SVG donut en el Hub principal con un score global de 0 a 100, donde 100 indica m\u00ednimo riesgo.", "Normal"),
        ]
        insert_multiple_after(anchor, new_66_content)
        print(f"  Secciones 6.6, 6.7, 6.8 insertadas despues de 6.5")

    # 7) Insert section 7.5 Traductor de Jerga Legal after MUPA
    idx_74, p_74 = find_paragraph(doc, "7.4 MUPA")
    if p_74:
        # Find end of 7.4 content
        next_idx = idx_74 + 1
        while next_idx < len(doc.paragraphs):
            next_p = doc.paragraphs[next_idx]
            next_text = next_p.text.strip()
            if next_text and (next_p.style.name.startswith('Heading 1') or next_text.startswith('8.')):
                break
            next_idx += 1
        anchor = doc.paragraphs[next_idx - 1]
        new_75_content = [
            ("7.5 Traductor de Jerga Legal", "Heading 2"),
            ("El Traductor de Jerga Legal es una herramienta impulsada por GPT-4o que traduce documentos legales, contratos y normativas del espa\u00f1ol jur\u00eddico formal a un lenguaje sencillo comprensible por emprendedores sin formaci\u00f3n legal.", "Normal"),
            ("El sistema utiliza un prompt especializado en derecho paname\u00f1o que estructura la respuesta en cuatro secciones: (1) EN SIMPLE \u2014 explicaci\u00f3n en lenguaje cotidiano, (2) PARA TU NEGOCIO \u2014 impacto pr\u00e1ctico en la operaci\u00f3n, (3) ACCI\u00d3N REQUERIDA \u2014 pasos concretos que el emprendedor debe tomar, y (4) DATOS IMPORTANTES \u2014 plazos, montos y referencias legales clave.", "Normal"),
            ("El usuario puede pegar texto legal completo o hacer preguntas sobre cl\u00e1usulas espec\u00edficas. La API route procesa el texto y retorna la traducci\u00f3n estructurada en formato card.", "Normal"),
        ]
        insert_multiple_after(anchor, new_75_content)
        print(f"  Seccion 7.5 Traductor de Jerga Legal insertada")

    # 8) Update section 8.1 capabilities - add Traductor Legal mention
    idx_81, p_81 = find_paragraph(doc, "8.1 Capacidades del Asistente")
    if p_81:
        # Find last list item in section 8.1
        idx_last = find_paragraph(doc, "Generaci\u00f3n autom\u00e1tica de asientos contables")[0]
        if idx_last:
            anchor = doc.paragraphs[idx_last]
            new_81_items = [
                ("Traducci\u00f3n de jerga legal a lenguaje sencillo con estructura EN SIMPLE / PARA TU NEGOCIO / ACCI\u00d3N REQUERIDA / DATOS IMPORTANTES", "List Paragraph"),
            ]
            insert_multiple_after(anchor, new_81_items)
            print(f"  Seccion 8.1: agregada capacidad de Traductor Legal")

    # 9) Insert section 9.1 Alertas de Facturacion after "9.1 Categor\u00edas de Alertas"
    # Actually the task says to add Facturacion as new alert category
    # Let's add after the last category in the alerts table
    idx_93, p_93 = find_paragraph(doc, "9.3 Comportamiento del")
    if p_93:
        # Insert before 9.3 (after 9.2 content)
        prev_p = doc.paragraphs[idx_93 - 1] if idx_93 > 0 else p_93
        new_alert_content = [
            ("9.1.1 Categor\u00eda de Facturaci\u00f3n (Nuevo)", "Heading 3"),
            ("Se agrega una nueva categor\u00eda de alertas para Facturaci\u00f3n que monitorea los umbrales de la Resoluci\u00f3n DGI 201-6299: alerta preventiva al alcanzar 80 facturas/mes (80% del l\u00edmite), alerta cr\u00edtica al superar 100 facturas/mes, alerta de tope de ingresos al acercarse a B/.36,000/a\u00f1o en ventas gravadas, y alerta de obligatoriedad de facturaci\u00f3n electr\u00f3nica v\u00eda PAC.", "Normal"),
        ]
        insert_multiple_after(prev_p, new_alert_content)
        print(f"  Seccion 9: agregada categoria de Facturacion en alertas")

    # 10) Insert section 10A Onboarding Guiado before Preguntas Frecuentes
    idx_10, p_10 = find_paragraph(doc, "10. Preguntas Frecuentes")
    if p_10:
        prev_p = doc.paragraphs[idx_10 - 1] if idx_10 > 0 else None
        if prev_p:
            new_10a_content = [
                ("10A. Onboarding Guiado", "Heading 1"),
                ("Al acceder por primera vez a la plataforma, el usuario es recibido por un tour modal interactivo de 5 pasos que presenta las funcionalidades principales de Mi Director Financiero PTY. El tour utiliza texto grande y alto contraste, optimizado para la Silver Economy (emprendedores mayores de 50 a\u00f1os).", "Normal"),
                ("Los 5 pasos del onboarding cubren: (1) Bienvenida y configuraci\u00f3n de empresa, (2) Diagn\u00f3stico Flash r\u00e1pido, (3) Navegaci\u00f3n por los tres m\u00f3dulos principales, (4) Sistema de alertas y notificaciones, (5) Mi Asistente IA y c\u00f3mo hacer consultas.", "Normal"),
                ("Adicionalmente, la plataforma incluye tooltips flotantes contextuales implementados con React createPortal, que aparecen junto a cada componente complejo para explicar su funci\u00f3n. Los tooltips se pueden desactivar desde la configuraci\u00f3n de usuario.", "Normal"),
            ]
            insert_multiple_after(prev_p, new_10a_content)
            print(f"  Seccion 10A Onboarding Guiado insertada")

    # 11) Update FAQ Q15 to mention CSV for ventas
    idx_15, p_15 = find_paragraph(doc, "15. \u00bfSe pueden importar datos")
    if p_15:
        # Find the answer paragraph (next non-empty one)
        ans_idx = idx_15 + 1
        if ans_idx < len(doc.paragraphs):
            ans_p = doc.paragraphs[ans_idx]
            # Replace the answer text
            for run in ans_p.runs:
                if "1.0.0" in run.text or "CSV" in run.text:
                    run.text = run.text.replace(
                        "La versi\u00f3n 1.1.0 soporta importaci\u00f3n de datos financieros mediante archivo CSV en el Diagn\u00f3stico Flash.",
                        "La versi\u00f3n 1.1.0 soporta importaci\u00f3n de datos financieros mediante archivo CSV en el Diagn\u00f3stico Flash, as\u00ed como importaci\u00f3n masiva de ventas desde archivos CSV del Sistema de Facturaci\u00f3n Electr\u00f3nica de Panam\u00e1 (SFEP) de la DGI en el Libro de Ventas."
                    )
            # Also try full text replacement if runs didn't match exactly
            full_text = ans_p.text
            if "Libro de Ventas" not in full_text and "CSV" in full_text:
                # Try to append info about ventas CSV at the end
                insert_paragraph_after(ans_p,
                    "Adicionalmente, el Libro de Ventas (v1.1.0) permite importar ventas masivas desde archivos CSV generados por el SFEP de la DGI, con validaci\u00f3n autom\u00e1tica de formato, RUC y montos.",
                    "Normal")
            print(f"  FAQ Q15: actualizada con mencion de CSV para ventas")

    # Save
    doc.save(str(MANUAL_OUT))
    print(f"\n  GUARDADO: {MANUAL_OUT}")


# ═══════════════════════════════════════════════════════════════
#  TASK 2: UPDATE MEMORIA DESCRIPTIVA
# ═══════════════════════════════════════════════════════════════

def update_memoria():
    print("\n" + "=" * 60)
    print("TASK 2: Actualizando Memoria Descriptiva...")
    print("=" * 60)

    doc = Document(str(MEMORIA_IN))

    # 1) Version 1.0.0 -> 1.1.0
    n = replace_text_in_all(doc, "1.0.0", "1.1.0")
    print(f"  Version: reemplazados {n} ocurrencias de 1.0.0 -> 1.1.0")

    # 2) Update metrics in table 3 (section 2.4)
    metrics_table = doc.tables[3]  # Table index 3 is the metrics table
    metric_updates = {
        "112 archivos": "185 archivos",
        "46,917 l\u00edneas": "76,000 l\u00edneas",
        "65 componentes": "95 componentes",
        "21 archivos": "35 archivos",
        "9 rutas": "19 rutas",
        "3 archivos SQL": "10 archivos SQL",
    }
    for row in metrics_table.rows:
        for cell in row.cells:
            for old_val, new_val in metric_updates.items():
                if old_val in cell.text:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            if old_val in run.text:
                                run.text = run.text.replace(old_val, new_val)
                    print(f"    Metrica actualizada: {old_val} -> {new_val}")

    # 3) Rename "Mandibulas de Cocodrilo" in section 3.2.3
    idx_mand, p_mand = find_paragraph(doc, "Mand\u00edbulas de Cocodrilo")
    if p_mand:
        for run in p_mand.runs:
            if "Mand\u00edbulas de Cocodrilo (Apalancamiento Operativo)" in run.text:
                run.text = run.text.replace(
                    "Mand\u00edbulas de Cocodrilo (Apalancamiento Operativo)",
                    "Brecha de Rentabilidad / Gr\u00e1fica de Tijeras (Apalancamiento Operativo)"
                )
            elif "Mand\u00edbulas de Cocodrilo" in run.text:
                run.text = run.text.replace(
                    "Mand\u00edbulas de Cocodrilo",
                    "Brecha de Rentabilidad"
                )
        print(f"  Renombrado: 3.2.3 Mandibulas -> Brecha de Rentabilidad")

    # Also rename in the paragraph content below and in tables
    for p in doc.paragraphs:
        if "Mand\u00edbulas" in p.text or "mand\u00edbulas" in p.text:
            for run in p.runs:
                # Replace full phrases first, then partial
                run.text = run.text.replace("Mand\u00edbulas de Cocodrilo", "Brecha de Rentabilidad (Gr\u00e1fica de Tijeras)")
                run.text = run.text.replace("mand\u00edbulas de cocodrilo", "brecha de rentabilidad")
                run.text = run.text.replace("Mand\u00edbulas", "Brecha de Rentabilidad")
                run.text = run.text.replace("mand\u00edbulas", "brecha de rentabilidad")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if "Mand\u00edbulas" in p.text or "mand\u00edbulas" in p.text:
                        for run in p.runs:
                            run.text = run.text.replace("Mand\u00edbulas de Cocodrilo", "Brecha de Rentabilidad (Gr\u00e1fica de Tijeras)")
                            run.text = run.text.replace("mand\u00edbulas de cocodrilo", "brecha de rentabilidad")
                            run.text = run.text.replace("Mand\u00edbulas", "Brecha de Rentabilidad")
                            run.text = run.text.replace("mand\u00edbulas", "brecha de rentabilidad")

    # Also in ASCII diagram
    # We'll handle the diagram update separately below

    # 4) Add new subsections in section 3.1 (Mi Contabilidad)
    # Insert 3.1.10 Libro de Ventas and 3.1.11 Cierre Mensual after 3.1.9
    idx_319, p_319 = find_paragraph(doc, "3.1.9 Mi Inventario")
    if p_319:
        # Find end of 3.1.9 content
        next_idx = idx_319 + 1
        while next_idx < len(doc.paragraphs):
            next_p = doc.paragraphs[next_idx]
            next_text = next_p.text.strip()
            if next_text and next_p.style.name.startswith('Heading 2'):
                break
            if next_text and next_p.style.name.startswith('Heading 3') and '3.2' in next_text:
                break
            next_idx += 1
        anchor = doc.paragraphs[next_idx - 1]

        new_31_content = [
            ("3.1.10 Libro de Ventas (3 Segmentos)", "Heading 3"),
            ("El Libro de Ventas implementa un sistema progresivo de registro de transacciones comerciales con tres segmentos de complejidad creciente, alineados con los umbrales de facturaci\u00f3n de la DGI establecidos en la Resoluci\u00f3n 201-6299.", "Normal"),
            ("Segmento 1 \u2014 Venta R\u00e1pida Manual: Bot\u00f3n de registro individual con c\u00e1lculo autom\u00e1tico de ITBMS (7% seg\u00fan Ley 8 de 2010). Badge gris en tabla unificada. Ideal para contribuyentes con bajo volumen.", "Normal"),
            ("Segmento 2 \u2014 Importaci\u00f3n CSV SFEP DGI: Parser que reconoce el formato oficial del Sistema de Facturaci\u00f3n Electr\u00f3nica de Panam\u00e1. Validaci\u00f3n autom\u00e1tica de RUC, montos e ITBMS. Badge azul en tabla unificada.", "Normal"),
            ("Segmento 3 \u2014 Integraci\u00f3n PAC Gosocket: Conexi\u00f3n con Proveedor Autorizado de Certificaci\u00f3n para facturaci\u00f3n electr\u00f3nica obligatoria cuando se superan 100 facturas/mes o B/.36,000/a\u00f1o. Badge verde en tabla unificada.", "Normal"),
            ("// Archivo: ventas-storage.ts", "Normal"),
            ("// Umbrales DGI Res. 201-6299:", "Normal"),
            ("const UMBRAL_FACTURAS_MES = 100;    // facturas/mes", "Normal"),
            ("const UMBRAL_VENTAS_ANUAL = 36000;  // B/.36,000/a\u00f1o", "Normal"),
            ("3.1.11 Cierre Mensual y Reportes", "Heading 3"),
            ("El m\u00f3dulo de Cierre Mensual genera reportes integrales del per\u00edodo fiscal que incluyen: snapshot de KPIs financieros (ingresos, EBITDA, margen neto, cash runway), cascada P&L del mes con comparativo vs. mes anterior, resumen de ventas por segmento de origen, estado del inventario, n\u00f3mina consolidada y alertas activas. El cierre protege los registros contra re-apertura, garantizando la integridad del historial contable.", "Normal"),
            ("// Archivo: reporte-engine.ts", "Normal"),
            ("// Genera snapshot mensual con proteccion contra re-apertura", "Normal"),
        ]
        insert_multiple_after(anchor, new_31_content)
        print(f"  Secciones 3.1.10 y 3.1.11 insertadas")

    # 5) Add new subsections in section 3.2 (Mis Finanzas)
    # Renumber existing sections? No, just add new ones at end
    # Find end of section 3.2 (before 3.3)
    idx_33, p_33 = find_paragraph(doc, "3.3 M\u00f3dulo: Mi Empresa")
    if p_33:
        anchor = doc.paragraphs[idx_33 - 1]
        new_32_content = [
            ("3.2.14 Forecast P&L 12 Meses", "Heading 3"),
            ("Proyecta el estado de resultados a 12 meses mediante regresi\u00f3n lineal client-side aplicada sobre los datos hist\u00f3ricos del negocio. Incluye cuatro supuestos editables (crecimiento de ingresos, incremento de costos, inflaci\u00f3n salarial, variaci\u00f3n OPEX), tabla P&L interactiva click-to-edit, y visualizaci\u00f3n con Recharts LineChart.", "Normal"),
            ("// Archivo: forecast-engine.ts", "Normal"),
            ("// Regresion lineal: y = mx + b", "Normal"),
            ("// m = (n\u03a3xy - \u03a3x\u03a3y) / (n\u03a3x\u00b2 - (\u03a3x)\u00b2)", "Normal"),
            ("// KPIs: revenue_anual, ebitda_estimado, margen_neto, meses_runway", "Normal"),
            ("3.2.15 Simulador de N\u00f3mina con Desglose Completo", "Heading 3"),
            ("Motor de c\u00e1lculo avanzado que presenta el desglose exhaustivo del costo laboral en Panam\u00e1. Vista individual con 14 l\u00edneas de desglose empleador (CSS_PATRONAL=12.25%, RP=1.50%, SE=1.50%, XIII Mes=8.33%, Vacaciones=3.33%, Antig\u00fcedad=1.92%) y c\u00e1lculo ISR paso a paso del empleado.", "Normal"),
            ("// Archivo: nomina-calculator.ts", "Normal"),
            ("// ISR verificado: B/.1,500 bruto", "Normal"),
            ("// Base gravable = 1500 \u00d7 0.89 = 1,335", "Normal"),
            ("// Anual = 1,335 \u00d7 12 = 16,020", "Normal"),
            ("// ISR anual = (16,020 - 11,000) \u00d7 0.15 = 753.00", "Normal"),
            ("// ISR mensual = 753.00 / 12 = 62.75", "Normal"),
            ("Vista consolidada multi-empleado con diferenciaci\u00f3n entre salida real de caja (salarios netos + cuotas obrero-patronales) y provisiones mensuales (XIII Mes, vacaciones, antig\u00fcedad).", "Normal"),
            ("3.2.16 Diagn\u00f3stico CB Insights (6 Causas de Fracaso)", "Heading 3"),
            ("Detector client-side que eval\u00faa las 6 causas principales de fracaso empresarial seg\u00fan el estudio de CB Insights, utilizando los datos financieros del usuario como insumo:", "Normal"),
            ("1. Sin mercado/demanda: evalua ingresos vs. break-even y tendencia.", "Normal"),
            ("2. Sin caja/liquidez: cash runway, prueba acida, CCC.", "Normal"),
            ("3. Equipo incorrecto: ratio nomina/ingresos, productividad.", "Normal"),
            ("4. Competencia superior: margen bruto vs. benchmark del rubro.", "Normal"),
            ("5. Precios inadecuados: margen de contribucion, elasticidad.", "Normal"),
            ("6. Modelo no escalable: apalancamiento operativo, tendencia EBITDA.", "Normal"),
            ("// Archivo: cb-insights-detector.ts", "Normal"),
            ("// Score global 0-100, widget SVG donut en Hub", "Normal"),
        ]
        insert_multiple_after(anchor, new_32_content)
        print(f"  Secciones 3.2.14, 3.2.15, 3.2.16 insertadas")

    # 6) Add new subsections in section 3.4 (Transversales)
    # Find end of section 3.4 (before section 4)
    idx_4, p_4 = find_paragraph(doc, "4. Aspectos Originales")
    if p_4:
        anchor = doc.paragraphs[idx_4 - 1]
        new_34_content = [
            ("3.4.5 Traductor de Jerga Legal", "Heading 3"),
            ("API route que utiliza GPT-4o con un system prompt especializado en derecho paname\u00f1o para traducir documentos legales, contratos y normativas a lenguaje sencillo. La respuesta se estructura en cuatro secciones: EN SIMPLE (explicaci\u00f3n cotidiana), PARA TU NEGOCIO (impacto pr\u00e1ctico), ACCI\u00d3N REQUERIDA (pasos concretos) y DATOS IMPORTANTES (plazos, montos, referencias).", "Normal"),
            ("// Archivo: api/herramientas/traductor-legal/route.ts", "Normal"),
            ("// System prompt: rol de abogado panameno, 4 secciones estructuradas", "Normal"),
            ("3.4.6 Onboarding Guiado", "Heading 3"),
            ("Tour modal interactivo de 5 pasos para nuevos usuarios, con texto grande y alto contraste optimizado para la Silver Economy. Implementado con React createPortal para tooltips flotantes contextuales que explican cada componente de la plataforma. Los pasos cubren: bienvenida, Diagn\u00f3stico Flash, navegaci\u00f3n por m\u00f3dulos, sistema de alertas, y Mi Asistente IA.", "Normal"),
            ("3.4.7 Panel de M\u00e9tricas Admin", "Heading 3"),
            ("Dashboard administrativo en /admin/metricas que presenta KPIs de uso de la plataforma: MAU (Monthly Active Users), churn rate mensual, ratio DAU/MAU (engagement), tabla de cohortes de retenci\u00f3n M1-M6, y gr\u00e1fico de barras de uso por m\u00f3dulo.", "Normal"),
        ]
        insert_multiple_after(anchor, new_34_content)
        print(f"  Secciones 3.4.5, 3.4.6, 3.4.7 insertadas")

    # 7) Add new differentiation items 16-20 in section 4
    # Note: the generic rename loop above may have already renamed this heading
    idx_15d, p_15d = find_paragraph(doc, "15. Brecha de Rentabilidad")
    if not p_15d:
        idx_15d, p_15d = find_paragraph(doc, "15. Mand\u00edbulas de Cocodrilo")
    if not p_15d:
        idx_15d, p_15d = find_paragraph(doc, "15. Mand")
    if p_15d:
        # Find end of item 15 content
        next_idx = idx_15d + 1
        while next_idx < len(doc.paragraphs):
            next_p = doc.paragraphs[next_idx]
            next_text = next_p.text.strip()
            if next_text and next_p.style.name.startswith('Heading'):
                break
            next_idx += 1
        anchor = doc.paragraphs[next_idx - 1]

        # Rename item 15 too
        for run in p_15d.runs:
            if "Mand\u00edbulas de Cocodrilo" in run.text:
                run.text = run.text.replace("Mand\u00edbulas de Cocodrilo", "Brecha de Rentabilidad (Gr\u00e1fica de Tijeras)")
            elif "Mand\u00edbulas" in run.text:
                run.text = run.text.replace("Mand\u00edbulas", "Brecha de Rentabilidad")

        new_items = [
            ("16. Libro de Ventas con 3 segmentos progresivos", "Heading 3"),
            ("Sistema \u00fanico en Panam\u00e1 que integra tres niveles de facturaci\u00f3n en una sola plataforma: venta r\u00e1pida manual para microempresas, importaci\u00f3n CSV desde el SFEP de la DGI para contribuyentes medianos, e integraci\u00f3n con PAC Gosocket para facturaci\u00f3n electr\u00f3nica obligatoria. Los umbrales de transici\u00f3n (100 facturas/mes, B/.36,000/a\u00f1o) son monitoreados autom\u00e1ticamente.", "Normal"),
            ("17. Forecast P&L con regresi\u00f3n lineal client-side", "Heading 3"),
            ("Proyecci\u00f3n financiera a 12 meses ejecutada \u00edntegramente en el navegador del usuario sin necesidad de conexi\u00f3n al servidor, utilizando regresi\u00f3n lineal sobre datos hist\u00f3ricos con cuatro supuestos editables y tabla P&L interactiva click-to-edit.", "Normal"),
            ("18. Detector CB Insights de 6 causas de fracaso", "Heading 3"),
            ("Implementaci\u00f3n original del framework de CB Insights adaptado al contexto paname\u00f1o, que eval\u00faa en tiempo real seis factores de riesgo empresarial utilizando los datos financieros del usuario y presenta un score visual SVG de 0-100.", "Normal"),
            ("19. Traductor de jerga legal con IA especializada", "Heading 3"),
            ("Herramienta sin equivalente en el mercado paname\u00f1o que traduce documentos legales a lenguaje sencillo mediante GPT-4o con prompt especializado en derecho paname\u00f1o, estructurando la respuesta en cuatro secciones accionables.", "Normal"),
            ("20. Simulador de n\u00f3mina con desglose de 14 l\u00edneas", "Heading 3"),
            ("Desglose exhaustivo del costo laboral en Panam\u00e1 que ning\u00fan otro software de gesti\u00f3n empresarial presenta con este nivel de detalle: 14 l\u00edneas del lado empleador con ISR paso a paso del empleado, verificado con las tasas vigentes de la Ley 462 de 2025.", "Normal"),
        ]
        insert_multiple_after(anchor, new_items)
        print(f"  Items 16-20 de diferenciacion agregados")

    # 8) Update ASCII diagram - replace the old one with updated version
    # Find the diagram start
    idx_diag, p_diag = find_paragraph(doc, "MI DIRECTOR FINANCIERO PTY")
    # Find the one inside section 5 that has the version
    for p in doc.paragraphs:
        if "MI DIRECTOR FINANCIERO PTY" in p.text and "Versi\u00f3n" in p.text:
            for run in p.runs:
                if "1.0.0" in run.text:
                    run.text = run.text.replace("1.0.0", "1.1.0")
            break

    # Update the diagram lines to include new modules
    # Find "Mand\u00edbulas de Cocodrilo" in diagram and replace
    for p in doc.paragraphs:
        if "Mand\u00edbulas de Cocodrilo" in p.text:
            for run in p.runs:
                if "Mand\u00edbulas de Cocodrilo" in run.text:
                    run.text = run.text.replace("Mand\u00edbulas de Cocodrilo", "Brecha de Rentabilidad  ")

    # Add new lines to the Mis Finanzas section of the diagram
    # Find line with "Ratios Financieros" in diagram
    idx_ratios = None
    for i, p in enumerate(doc.paragraphs):
        if "Ratios Financieros" in p.text and "|" in p.text:
            idx_ratios = i
            break

    if idx_ratios:
        anchor = doc.paragraphs[idx_ratios]
        new_diag_lines = [
            ("| |   Forecast P&L 12 Meses        |  |                                         | |", "Normal"),
            ("| |   Simulador N\u00f3mina Desglose    |  |                                         | |", "Normal"),
            ("| |   Diagn\u00f3stico CB Insights      |  |                                         | |", "Normal"),
        ]
        insert_multiple_after(anchor, new_diag_lines)
        print(f"  Diagrama ASCII: agregadas lineas de nuevos modulos financieros")

    # Add new lines to Mi Contabilidad section of the diagram
    # Find line with "Costo Prom. Ponderado" in diagram
    idx_costo = None
    for i, p in enumerate(doc.paragraphs):
        if "Costo Prom. Ponderado" in p.text and "|" in p.text:
            idx_costo = i
            break

    if idx_costo:
        anchor = doc.paragraphs[idx_costo]
        new_diag_lines_cont = [
            ("| | Libro de Ventas (3 seg.)  |  |                                         | |", "Normal"),
            ("| | Cierre Mensual + Reportes |  |                                         | |", "Normal"),
        ]
        insert_multiple_after(anchor, new_diag_lines_cont)
        print(f"  Diagrama ASCII: agregadas lineas Libro Ventas y Cierre Mensual")

    # Update the transversal functions line to include new items
    idx_trans = None
    for i, p in enumerate(doc.paragraphs):
        if "Glosario Inteligente" in p.text and "Facturaci\u00f3n" in p.text:
            idx_trans = i
            break

    if idx_trans:
        anchor = doc.paragraphs[idx_trans]
        new_trans_lines = [
            ("| | Traductor Legal (GPT-4o) | Onboarding Guiado (5 pasos)               | |", "Normal"),
            ("| | CB Insights Detector (6 causas) | Forecast P&L | Admin Metricas      | |", "Normal"),
        ]
        insert_multiple_after(anchor, new_trans_lines)
        print(f"  Diagrama ASCII: agregadas lineas transversales")

    # 9) Add new row to Historial de Desarrollo table (table 6)
    hist_table = doc.tables[6]
    new_row = hist_table.add_row()
    cells = new_row.cells
    cells[0].text = "7 mar 2026"
    cells[1].text = "v1.1.0"
    cells[2].text = (
        "Libro de Ventas (3 segmentos), Forecast P&L 12 meses, "
        "Simulador N\u00f3mina con desglose, Diagn\u00f3stico CB Insights, "
        "Traductor Legal (GPT-4o), Cierre Mensual, Onboarding Guiado, "
        "Panel Admin M\u00e9tricas, alertas de Facturaci\u00f3n. "
        "Branding: Mand\u00edbulas renombrado a Brecha de Rentabilidad. "
        "~185 archivos, ~76,000 l\u00edneas de c\u00f3digo."
    )
    print(f"  Historial: agregada fila para 7 mar 2026 v1.1.0")

    # Also update the item count text in section 4 header
    idx_4h, p_4h = find_paragraph(doc, "quince elementos de originalidad")
    if p_4h:
        for run in p_4h.runs:
            if "quince" in run.text:
                run.text = run.text.replace("quince", "veinte")
        print(f"  Seccion 4: 'quince' -> 'veinte' elementos de originalidad")

    # Save
    doc.save(str(MEMORIA_OUT))
    print(f"\n  GUARDADO: {MEMORIA_OUT}")


# ═══════════════════════════════════════════════════════════════
#  TASK 3: GENERATE HTML CODIGO FUENTE
# ═══════════════════════════════════════════════════════════════

SOURCE_FILES = [
    "frontend/src/app/dashboard/page.tsx",
    "frontend/src/lib/calculations.ts",
    "frontend/src/lib/alerts.ts",
    "frontend/src/components/rrhh/MiRRHH.tsx",
    "frontend/src/lib/ventas-storage.ts",
    "frontend/src/lib/analytics/cb-insights-detector.ts",
    "frontend/src/lib/rrhh/nomina-calculator.ts",
    "frontend/src/lib/analytics/forecast-engine.ts",
    "frontend/src/services/facturacion/FacturacionService.ts",
    "frontend/src/lib/reportes/reporte-engine.ts",
]

FIRST_N_LINES = 80
LAST_N_LINES = 50


def read_file_lines(filepath):
    """Lee un archivo y retorna sus lineas."""
    full_path = REPO_ROOT / filepath
    if not full_path.exists():
        return [f"// ARCHIVO NO ENCONTRADO: {filepath}"]
    with open(full_path, "r", encoding="utf-8", errors="replace") as f:
        return f.readlines()


def generate_html():
    print("\n" + "=" * 60)
    print("TASK 3: Generando HTML de Codigo Fuente...")
    print("=" * 60)

    # Collect file data
    file_data = []
    for rel_path in SOURCE_FILES:
        lines = read_file_lines(rel_path)
        total = len(lines)
        first = lines[:FIRST_N_LINES]
        last = lines[-LAST_N_LINES:] if total > LAST_N_LINES else lines
        file_data.append({
            "path": rel_path,
            "total_lines": total,
            "first_lines": first,
            "last_lines": last,
        })

    # Build HTML
    html_parts = []

    # CSS and page setup
    html_parts.append("""<!DOCTYPE html>
<html lang="es">
<head>
<meta charset="UTF-8">
<title>C\u00f3digo Fuente - Mi Director Financiero PTY</title>
<style>
@page {
    size: letter;
    margin: 2cm 2.5cm 2.5cm 2.5cm;
    @bottom-center {
        content: "P\u00e1gina " counter(page) " de " counter(pages);
        font-family: 'Courier New', monospace;
        font-size: 9pt;
        color: #666;
    }
}

* {
    margin: 0;
    padding: 0;
    box-sizing: border-box;
}

body {
    font-family: 'Courier New', Courier, monospace;
    font-size: 9pt;
    line-height: 1.4;
    color: #1a1a1a;
    background: white;
}

.cover-page {
    page-break-after: always;
    display: flex;
    flex-direction: column;
    justify-content: center;
    align-items: center;
    min-height: 90vh;
    text-align: center;
    padding: 3cm;
}

.cover-page h1 {
    font-family: 'Georgia', 'Times New Roman', serif;
    font-size: 28pt;
    font-weight: bold;
    letter-spacing: 4pt;
    margin-bottom: 1cm;
    color: #1A242F;
}

.cover-page h2 {
    font-family: 'Georgia', 'Times New Roman', serif;
    font-size: 18pt;
    font-weight: normal;
    margin-bottom: 0.5cm;
    color: #C5A059;
}

.cover-page .subtitle {
    font-size: 12pt;
    margin-bottom: 0.3cm;
    color: #555;
}

.cover-page .version {
    font-size: 14pt;
    font-weight: bold;
    margin: 0.8cm 0;
    color: #1A242F;
}

.cover-page .author {
    font-size: 11pt;
    margin-top: 1.5cm;
    color: #333;
    line-height: 1.8;
}

.cover-page .footer-info {
    font-size: 10pt;
    margin-top: 2cm;
    color: #777;
    line-height: 1.8;
    border-top: 1px solid #ccc;
    padding-top: 0.5cm;
}

.toc-page {
    page-break-after: always;
    padding: 1cm 0;
}

.toc-page h2 {
    font-family: 'Georgia', 'Times New Roman', serif;
    font-size: 16pt;
    margin-bottom: 1cm;
    color: #1A242F;
    border-bottom: 2px solid #C5A059;
    padding-bottom: 0.3cm;
}

.toc-page .toc-section {
    font-size: 10pt;
    font-weight: bold;
    margin-top: 0.5cm;
    margin-bottom: 0.2cm;
    color: #1A242F;
}

.toc-page .toc-item {
    font-size: 9pt;
    margin-left: 1cm;
    margin-bottom: 0.1cm;
    color: #555;
}

.file-section {
    page-break-before: always;
}

.file-header {
    background: #1A242F;
    color: #C5A059;
    padding: 8pt 12pt;
    font-family: 'Courier New', monospace;
    font-size: 10pt;
    font-weight: bold;
    margin-bottom: 0;
    border-radius: 4pt 4pt 0 0;
}

.file-meta {
    background: #f0f0f0;
    padding: 4pt 12pt;
    font-size: 8pt;
    color: #666;
    border-bottom: 1px solid #ddd;
}

.section-label {
    background: #C5A059;
    color: white;
    padding: 4pt 12pt;
    font-size: 9pt;
    font-weight: bold;
    margin-top: 12pt;
}

pre {
    background: #fafafa;
    border: 1px solid #e0e0e0;
    border-top: none;
    padding: 8pt 12pt;
    font-family: 'Courier New', Courier, monospace;
    font-size: 8pt;
    line-height: 1.35;
    overflow-wrap: break-word;
    white-space: pre-wrap;
    tab-size: 2;
}

.line-num {
    color: #999;
    display: inline-block;
    width: 4ch;
    text-align: right;
    margin-right: 1.5ch;
    user-select: none;
}

.ellipsis {
    text-align: center;
    padding: 6pt;
    color: #999;
    font-style: italic;
    font-size: 9pt;
    background: #f5f5f5;
    border: 1px dashed #ddd;
}

@media print {
    body { font-size: 8pt; }
    pre { font-size: 7.5pt; }
    .file-section { page-break-before: always; }
    .cover-page { page-break-after: always; }
    .toc-page { page-break-after: always; }
}
</style>
</head>
<body>
""")

    # Cover page
    html_parts.append("""
<div class="cover-page">
    <h1>C\u00d3DIGO FUENTE</h1>
    <h2>MI DIRECTOR FINANCIERO PTY</h2>
    <div class="subtitle">Programa de Ordenador (Software)</div>
    <div class="version">Versi\u00f3n 1.1.0</div>

    <div class="author">
        Autora: Soraya Dayan Guti\u00e9rrez P\u00e9rez<br>
        Titular: Soraya Dayan Guti\u00e9rrez P\u00e9rez (Persona Natural)
    </div>

    <div class="subtitle" style="margin-top: 1.5cm;">
        Panam\u00e1, marzo de 2026
    </div>

    <div class="footer-info">
        Registro ante la Direcci\u00f3n Nacional de Derecho de Autor<br>
        Ministerio de Cultura &mdash; Rep\u00fablica de Panam\u00e1
    </div>
</div>
""")

    # Table of Contents
    html_parts.append("""
<div class="toc-page">
    <h2>TABLA DE CONTENIDO</h2>
    <div class="toc-section">Archivos incluidos (primeras y \u00faltimas l\u00edneas de cada uno):</div>
""")

    for i, fd in enumerate(file_data, 1):
        html_parts.append(f'    <div class="toc-item">{i}. {html_module.escape(fd["path"])} ({fd["total_lines"]} l\u00edneas)</div>\n')

    html_parts.append("""
    <div class="toc-section" style="margin-top: 1cm;">Nota:</div>
    <div class="toc-item">
        Para cada archivo se muestran las primeras ~80 l\u00edneas (equivalente a ~3 p\u00e1ginas)<br>
        y las \u00faltimas ~50 l\u00edneas (equivalente a ~2 p\u00e1ginas) del c\u00f3digo fuente original.
    </div>
    <div class="toc-item" style="margin-top: 0.3cm;">
        Total de archivos de c\u00f3digo fuente en el proyecto: ~185 archivos<br>
        Total de l\u00edneas de c\u00f3digo aproximadas: ~76,000 l\u00edneas
    </div>
</div>
""")

    # File sections
    for fd in file_data:
        path = fd["path"]
        total = fd["total_lines"]
        first = fd["first_lines"]
        last = fd["last_lines"]

        html_parts.append(f'<div class="file-section">\n')
        html_parts.append(f'<div class="file-header">{html_module.escape(path)}</div>\n')
        html_parts.append(f'<div class="file-meta">Total: {total} l\u00edneas | Mostrando primeras {len(first)} y \u00faltimas {len(last)} l\u00edneas</div>\n')

        # First lines
        html_parts.append(f'<div class="section-label">PRIMERAS {len(first)} L\u00cdNEAS</div>\n')
        html_parts.append('<pre>')
        for j, line in enumerate(first, 1):
            escaped = html_module.escape(line.rstrip('\n\r'))
            html_parts.append(f'<span class="line-num">{j}</span>{escaped}\n')
        html_parts.append('</pre>\n')

        # Ellipsis
        if total > FIRST_N_LINES + LAST_N_LINES:
            omitted = total - FIRST_N_LINES - LAST_N_LINES
            html_parts.append(f'<div class="ellipsis">... {omitted} l\u00edneas omitidas ...</div>\n')

        # Last lines
        start_line = total - len(last) + 1
        html_parts.append(f'<div class="section-label">\u00daLTIMAS {len(last)} L\u00cdNEAS</div>\n')
        html_parts.append('<pre>')
        for j, line in enumerate(last):
            line_num = start_line + j
            escaped = html_module.escape(line.rstrip('\n\r'))
            html_parts.append(f'<span class="line-num">{line_num}</span>{escaped}\n')
        html_parts.append('</pre>\n')

        html_parts.append('</div>\n\n')

    html_parts.append("""
</body>
</html>
""")

    # Write HTML
    os.makedirs(HTML_OUT.parent, exist_ok=True)
    with open(HTML_OUT, "w", encoding="utf-8") as f:
        f.write("".join(html_parts))

    print(f"\n  GUARDADO: {HTML_OUT}")


# ═══════════════════════════════════════════════════════════════
#  MAIN
# ═══════════════════════════════════════════════════════════════

def main():
    print("\n" + "=" * 60)
    print("  DNDA Documents Updater - Mi Director Financiero PTY")
    print("  Version 1.0.0 -> 1.1.0")
    print("=" * 60)

    # Verify inputs exist
    for path in [MANUAL_IN, MEMORIA_IN]:
        if not path.exists():
            print(f"ERROR: No se encontro {path}")
            sys.exit(1)

    # Verify repo exists
    if not REPO_ROOT.exists():
        print(f"ERROR: No se encontro el repositorio en {REPO_ROOT}")
        sys.exit(1)

    # Run tasks
    update_manual()
    update_memoria()
    generate_html()

    print("\n" + "=" * 60)
    print("  COMPLETADO")
    print("=" * 60)
    print(f"\n  Archivos generados:")
    print(f"    1. {MANUAL_OUT}")
    print(f"    2. {MEMORIA_OUT}")
    print(f"    3. {HTML_OUT}")
    print()


if __name__ == "__main__":
    main()
